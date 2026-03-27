"""Document parser service for extracting text from uploaded evidence files.

Supports: PDF, DOCX, XLSX/CSV, JSON, XML, plain text, and images.
Falls back gracefully if optional parsing libraries are not installed.
"""
import io
import json
import csv
import re
import logging
from typing import Optional, Dict, Any, Tuple

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse uploaded documents and extract text content + metadata."""

    @staticmethod
    def parse(file_bytes: bytes, file_name: str, mime_type: str) -> dict:
        """
        Parse a document and return extracted text + metadata.
        
        Returns:
            {"text": str, "metadata": dict}
        """
        file_name_lower = file_name.lower()
        
        try:
            # Route by mime type / extension
            if mime_type == "application/pdf" or file_name_lower.endswith(".pdf"):
                return DocumentParser._parse_pdf(file_bytes, file_name)
            
            elif mime_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ) or file_name_lower.endswith(".docx"):
                return DocumentParser._parse_docx(file_bytes, file_name)
            
            elif mime_type in (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            ) or file_name_lower.endswith((".xlsx", ".xls")):
                return DocumentParser._parse_xlsx(file_bytes, file_name)
            
            elif mime_type == "text/csv" or file_name_lower.endswith(".csv"):
                return DocumentParser._parse_csv(file_bytes, file_name)
            
            elif mime_type == "application/json" or file_name_lower.endswith(".json"):
                return DocumentParser._parse_json(file_bytes, file_name)
            
            elif file_name_lower.endswith(".nessus"):
                # Preserve raw XML for structured CVE extraction by intelligence_service
                return DocumentParser._parse_nessus(file_bytes, file_name)

            elif mime_type in ("application/xml", "text/xml") or file_name_lower.endswith(".xml"):
                return DocumentParser._parse_xml(file_bytes, file_name)
            
            elif mime_type.startswith("text/") or file_name_lower.endswith((".txt", ".md", ".log", ".yml", ".yaml", ".ini", ".cfg", ".conf")):
                return DocumentParser._parse_text(file_bytes, file_name)
            
            elif mime_type.startswith("image/"):
                return DocumentParser._parse_image(file_bytes, file_name)
            
            else:
                # Try as text, fall back to binary info
                return DocumentParser._parse_text(file_bytes, file_name)

        except Exception as e:
            logger.error(f"Error parsing {file_name} ({mime_type}): {e}")
            return {
                "text": f"[Parsing failed for {file_name}: {str(e)}]",
                "metadata": {
                    "parser": "error",
                    "error": str(e),
                    "file_name": file_name,
                    "mime_type": mime_type,
                    "size_bytes": len(file_bytes)
                }
            }

    @staticmethod
    def _parse_pdf(file_bytes: bytes, file_name: str) -> dict:
        """Extract text from PDF files. Tries PyPDF2 first, then pdfplumber."""
        # Try PyPDF2 first (lightweight, always available in Lambda)
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                except Exception as page_err:
                    logger.warning(f"PyPDF2 failed on page {i} of {file_name}: {page_err}")
                    text = ""
                if text.strip():
                    pages.append(text)

            full_text = "\n\n--- Page Break ---\n\n".join(pages)
            if full_text.strip():
                return {
                    "text": full_text,
                    "metadata": {
                        "parser": "PyPDF2",
                        "page_count": len(reader.pages),
                        "pages_with_text": len(pages),
                        "file_name": file_name,
                        "char_count": len(full_text)
                    }
                }
            # PyPDF2 extracted no text — try pdfplumber as fallback
            logger.info(f"PyPDF2 extracted no text from {file_name}, trying pdfplumber")
        except ImportError:
            logger.info("PyPDF2 not available, trying pdfplumber")
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_name}: {e}, trying pdfplumber")

        # Fallback: try pdfplumber (better for scanned/complex PDFs)
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text() or ""
                        tables = page.extract_tables()
                        table_text = ""
                        for table in tables:
                            for row in table:
                                table_text += " | ".join(str(cell or "") for cell in row) + "\n"
                        combined = text
                        if table_text:
                            combined += f"\n[Table Data]\n{table_text}"
                        if combined.strip():
                            pages.append(combined)
                    except Exception as page_err:
                        logger.warning(f"pdfplumber failed on page {i} of {file_name}: {page_err}")

            full_text = "\n\n--- Page Break ---\n\n".join(pages)
            return {
                "text": full_text or f"[PDF file: {file_name} - no extractable text found]",
                "metadata": {
                    "parser": "pdfplumber",
                    "page_count": len(pages),
                    "file_name": file_name,
                    "char_count": len(full_text)
                }
            }
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"pdfplumber also failed for {file_name}: {e}")

        # Both parsers unavailable or failed
        return {
            "text": f"[PDF file: {file_name} - could not extract text. Size: {len(file_bytes)} bytes]",
            "metadata": {"parser": "none", "file_name": file_name, "size_bytes": len(file_bytes)}
        }

    @staticmethod
    def _parse_pdf_via_multimodal(file_bytes: bytes, file_name: str) -> Optional[dict]:
        """
        Fallback for scanned PDFs: send pages as images to Bedrock multimodal.
        Only called when text-based extraction yields <100 chars.
        """
        from ..core.config import settings

        ocr_enabled = getattr(settings, "ocr_via_bedrock", True) and getattr(settings, "bedrock_enabled", False)
        if not ocr_enabled:
            return None

        try:
            from ..services.bedrock_service import bedrock_service

            # Try to render first 3 pages as images using PyPDF2 + PIL
            # PyPDF2 doesn't render pages — we'd need pdf2image which requires poppler.
            # Instead, send the raw PDF bytes info as context and let bedrock analyze
            # the text we do have. For true scanned-PDF OCR we'd need a rendering lib.
            # This fallback informs the AI about what we're dealing with.
            logger.info(f"PDF {file_name} appears to be scanned — Bedrock multimodal PDF OCR not available without pdf2image")
            return None
        except Exception as e:
            logger.debug(f"Scanned PDF multimodal fallback failed for {file_name}: {e}")
            return None

    @staticmethod
    def _parse_docx(file_bytes: bytes, file_name: str) -> dict:
        """Extract text from DOCX files."""
        try:
            from docx import Document
        except ImportError:
            return {
                "text": f"[DOCX file: {file_name} - python-docx not installed]",
                "metadata": {"parser": "none", "file_name": file_name, "size_bytes": len(file_bytes)}
            }

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading structure
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        prefix = "#" * int(level)
                    except ValueError:
                        prefix = "#"
                    paragraphs.append(f"{prefix} {para.text}")
                else:
                    paragraphs.append(para.text)

        # Also extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_texts.append("\n".join(rows))

        full_text = "\n\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n[Tables]\n" + "\n\n".join(table_texts)

        return {
            "text": full_text,
            "metadata": {
                "parser": "python-docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "file_name": file_name,
                "char_count": len(full_text)
            }
        }

    @staticmethod
    def _parse_xlsx(file_bytes: bytes, file_name: str) -> dict:
        """Extract text from Excel files."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {
                "text": f"[Excel file: {file_name} - openpyxl not installed]",
                "metadata": {"parser": "none", "file_name": file_name, "size_bytes": len(file_bytes)}
            }

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        sheets_text = []
        total_rows = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cell_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in cell_values):
                    rows.append(" | ".join(cell_values))
            
            if rows:
                total_rows += len(rows)
                sheet_text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
                sheets_text.append(sheet_text)

        wb.close()
        full_text = "\n\n".join(sheets_text)

        return {
            "text": full_text,
            "metadata": {
                "parser": "openpyxl",
                "sheet_count": len(wb.sheetnames),
                "total_rows": total_rows,
                "file_name": file_name,
                "char_count": len(full_text)
            }
        }

    @staticmethod
    def _parse_csv(file_bytes: bytes, file_name: str) -> dict:
        """Extract text from CSV files."""
        text = file_bytes.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = []
        for row in reader:
            rows.append(" | ".join(row))

        full_text = "\n".join(rows)
        return {
            "text": full_text,
            "metadata": {
                "parser": "csv",
                "row_count": len(rows),
                "file_name": file_name,
                "char_count": len(full_text)
            }
        }

    @staticmethod
    def _parse_json(file_bytes: bytes, file_name: str) -> dict:
        """Extract structured text from JSON files."""
        text = file_bytes.decode("utf-8", errors="replace")
        
        try:
            data = json.loads(text)
            # Pretty print for readability
            formatted = json.dumps(data, indent=2, default=str)
            
            # For large JSON, summarize structure
            metadata = {
                "parser": "json",
                "file_name": file_name,
                "char_count": len(formatted)
            }
            
            if isinstance(data, list):
                metadata["record_count"] = len(data)
                metadata["type"] = "array"
            elif isinstance(data, dict):
                metadata["top_level_keys"] = list(data.keys())[:20]
                metadata["type"] = "object"
            
            # Truncate very large JSON to keep within token limits
            if len(formatted) > 50000:
                formatted = formatted[:50000] + "\n... [truncated, full file has " + str(len(text)) + " chars]"
            
            return {"text": formatted, "metadata": metadata}

        except json.JSONDecodeError:
            return {"text": text[:50000], "metadata": {"parser": "json_raw", "file_name": file_name}}

    @staticmethod
    def _parse_nessus(file_bytes: bytes, file_name: str) -> dict:
        """
        Store the raw Nessus XML so intelligence_service can parse CVE structure.

        The generic _parse_xml flattens elements into tag:value text which loses
        the <cve>, severity attributes, and pluginID grouping that intelligence_service
        needs.  We preserve the raw XML here (up to 500 KB) and fall back to
        flattened text only if the file is too large.
        """
        MAX_RAW_XML_BYTES = 500_000

        text = file_bytes.decode("utf-8", errors="replace")
        element_count = 0

        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(text)
            element_count = sum(1 for _ in root.iter())
        except Exception:
            pass  # store raw text anyway; intelligence_service will handle parse errors

        if len(text) <= MAX_RAW_XML_BYTES:
            stored_text = text
        else:
            # Truncation could break XML — store flat fallback instead
            texts = []
            try:
                import xml.etree.ElementTree as ET
                root = ET.fromstring(text)
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                        texts.append(f"{tag}: {elem.text.strip()}")
                    if elem.attrib:
                        for k, v in elem.attrib.items():
                            texts.append(f"  @{k}={v}")
            except Exception:
                pass
            stored_text = "\n".join(texts)[:MAX_RAW_XML_BYTES]

        return {
            "text": stored_text,
            "metadata": {
                "parser": "nessus_xml",
                "file_name": file_name,
                "size_bytes": len(file_bytes),
                "char_count": len(stored_text),
                "element_count": element_count,
                "raw_xml_preserved": len(text) <= MAX_RAW_XML_BYTES,
            },
        }

    @staticmethod
    def _parse_xml(file_bytes: bytes, file_name: str) -> dict:
        """Extract text from XML files."""
        text = file_bytes.decode("utf-8", errors="replace")
        
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(text)
            
            # Extract all text content
            texts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                    texts.append(f"{tag}: {elem.text.strip()}")
                if elem.attrib:
                    for k, v in elem.attrib.items():
                        texts.append(f"  @{k}={v}")

            full_text = "\n".join(texts)
            return {
                "text": full_text if full_text else text[:50000],
                "metadata": {
                    "parser": "xml",
                    "root_tag": root.tag,
                    "element_count": sum(1 for _ in root.iter()),
                    "file_name": file_name,
                    "char_count": len(full_text)
                }
            }
        except Exception:
            # Fall back to raw XML text
            return {
                "text": text[:50000],
                "metadata": {"parser": "xml_raw", "file_name": file_name, "char_count": len(text)}
            }

    @staticmethod
    def _parse_text(file_bytes: bytes, file_name: str) -> dict:
        """Extract text from plain text files."""
        try:
            text = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            text = file_bytes.decode("latin-1", errors="replace")

        if len(text) > 50000:
            text = text[:50000] + f"\n... [truncated, full file has {len(file_bytes)} bytes]"

        return {
            "text": text,
            "metadata": {
                "parser": "text",
                "file_name": file_name,
                "char_count": len(text)
            }
        }

    @staticmethod
    def _parse_image(file_bytes: bytes, file_name: str) -> dict:
        """Handle image files — use Bedrock multimodal for content extraction when available."""
        from ..core.config import settings

        metadata = {
            "parser": "image",
            "file_name": file_name,
            "size_bytes": len(file_bytes)
        }

        # Capture basic image metadata from PIL
        mime_type = "image/png"
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(file_bytes))
            metadata["width"] = img.width
            metadata["height"] = img.height
            metadata["format"] = img.format
            metadata["mode"] = img.mode
            # Map PIL format to MIME type
            fmt_map = {"PNG": "image/png", "JPEG": "image/jpeg", "GIF": "image/gif", "WEBP": "image/webp"}
            mime_type = fmt_map.get(img.format, "image/png")
        except ImportError:
            # Infer MIME type from filename
            ext = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else "png"
            mime_type = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "webp": "image/webp"}.get(ext, "image/png")
        except Exception as e:
            metadata["error"] = str(e)

        # Use Bedrock multimodal if enabled and image is within size limit
        ocr_enabled = getattr(settings, "ocr_via_bedrock", True) and getattr(settings, "bedrock_enabled", False)
        max_bytes = getattr(settings, "ocr_max_image_size_mb", 5) * 1024 * 1024

        if ocr_enabled and len(file_bytes) <= max_bytes:
            try:
                from ..services.bedrock_service import bedrock_service

                prompt = (
                    "Analyze this image as a cybersecurity professional. "
                    "Extract ALL visible text, labels, IP addresses, hostnames, and annotations. "
                    "If this is a network diagram or architecture diagram, describe the complete topology: "
                    "identify all components (servers, databases, firewalls, load balancers, subnets, VPNs), "
                    "their connections, trust boundaries, and any security-relevant details. "
                    "If this is a screenshot of a scan report or configuration, extract all readable content. "
                    "Be thorough — include every piece of text and structural detail visible."
                )

                description = bedrock_service.analyze_image(
                    image_bytes=file_bytes,
                    mime_type=mime_type,
                    prompt=prompt,
                )

                if description and len(description.strip()) > 20:
                    metadata["parser"] = "bedrock_multimodal"
                    metadata["char_count"] = len(description)
                    return {
                        "text": description,
                        "metadata": metadata,
                    }

                logger.info(f"Bedrock multimodal returned insufficient content for {file_name}")
            except Exception as e:
                logger.warning(f"Bedrock multimodal failed for {file_name}: {e}")

        # Fallback: return PIL metadata only
        return {
            "text": f"[Image file: {file_name}, size: {len(file_bytes)} bytes. Image content cannot be parsed as text. Consider using this as a visual reference.]",
            "metadata": metadata
        }

    @staticmethod
    def detect_document_type(text: str, file_name: str) -> str:
        """
        Auto-detect the document type based on content and filename.
        Returns: vulnerability_scan, architecture_doc, policy, network_diagram, config, other
        """
        doc_type, _ = DocumentParser.detect_document_type_with_confidence(text, file_name)
        return doc_type

    @staticmethod
    def detect_document_type_with_confidence(text: str, file_name: str) -> Tuple[str, int]:
        """
        Auto-detect document type with a confidence score (0-100).
        Returns: (document_type, confidence)
        """
        text_lower = (text or "").lower()
        name_lower = file_name.lower()

        scores: Dict[str, int] = {
            "vulnerability_scan": 0,
            "architecture_doc": 0,
            "policy": 0,
            "config": 0,
            "network_diagram": 0,
        }

        # Vulnerability scan indicators
        vuln_keywords = ["cve-", "cvss", "vulnerability", "exploit", "severity", "critical", "nessus",
                        "qualys", "openvas", "nmap", "port scan", "finding", "remediation"]
        scores["vulnerability_scan"] = sum(1 for kw in vuln_keywords if kw in text_lower)
        if "scan" in name_lower or "vuln" in name_lower or "nessus" in name_lower:
            scores["vulnerability_scan"] += 3

        # Architecture document indicators
        arch_keywords = ["architecture", "component", "microservice", "api gateway", "load balancer",
                        "database", "deployment", "infrastructure", "topology", "data flow",
                        "system design", "high level design", "hld", "lld", "detailed design"]
        scores["architecture_doc"] = sum(1 for kw in arch_keywords if kw in text_lower)
        if "architect" in name_lower or "design" in name_lower or "hld" in name_lower:
            scores["architecture_doc"] += 3

        # Network diagram indicators (text extracted from diagrams via Bedrock multimodal)
        net_keywords = ["subnet", "vlan", "dmz", "firewall", "router", "switch",
                       "10.", "192.168.", "172.16.", "cidr", "/24", "/16",
                       "network diagram", "network topology", "ingress", "egress"]
        scores["network_diagram"] = sum(1 for kw in net_keywords if kw in text_lower)
        if "network" in name_lower or "diagram" in name_lower or "topology" in name_lower:
            scores["network_diagram"] += 3

        # Policy document indicators
        policy_keywords = ["policy", "compliance", "regulation", "standard", "requirement",
                          "control", "iso 27001", "nist", "soc 2", "gdpr", "hipaa", "pci"]
        scores["policy"] = sum(1 for kw in policy_keywords if kw in text_lower)
        if "policy" in name_lower or "compliance" in name_lower:
            scores["policy"] += 3

        # Configuration indicators
        config_keywords = ["server", "port", "host", "config", "setting", "parameter",
                          "environment", "variable", "connection string"]
        scores["config"] = sum(1 for kw in config_keywords if kw in text_lower)
        if "config" in name_lower or name_lower.endswith((".yml", ".yaml", ".ini", ".cfg", ".conf", ".env")):
            scores["config"] += 3

        # Find the highest-scoring type
        best_type = max(scores, key=lambda k: scores[k])
        best_score = scores[best_type]

        if best_score < 3:
            return ("other", max(10, min(30, best_score * 15)))

        # Confidence: scale the raw keyword count to 0-100
        confidence = min(95, 40 + best_score * 8)
        return (best_type, confidence)

    @staticmethod
    def extract_structured_metadata(text: str, document_type: str, file_name: str) -> Dict[str, Any]:
        """
        Extract structured metadata from document text based on document type.
        Runs BEFORE Bedrock analysis to provide quick stats.
        """
        metadata: Dict[str, Any] = {}

        if document_type == "vulnerability_scan":
            metadata = DocumentParser._extract_vuln_scan_metadata(text)
        elif document_type == "config":
            metadata = DocumentParser._extract_config_metadata(text)
        elif document_type == "policy":
            metadata = DocumentParser._extract_policy_metadata(text)
        elif document_type in ("architecture_doc", "network_diagram"):
            metadata = DocumentParser._extract_architecture_metadata(text)

        return metadata

    @staticmethod
    def _extract_vuln_scan_metadata(text: str) -> Dict[str, Any]:
        """Extract CVE IDs, severity counts, unique hosts from vulnerability scan text."""
        cve_pattern = re.compile(r'CVE-\d{4}-\d{4,7}', re.IGNORECASE)
        cve_ids = list(dict.fromkeys(cve_pattern.findall(text)))

        text_lower = text.lower()
        severity_counts = {
            "critical": text_lower.count("critical"),
            "high": text_lower.count("high"),
            "medium": text_lower.count("medium"),
            "low": text_lower.count("low"),
        }

        # Detect hosts (IP addresses)
        ip_pattern = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')
        hosts = list(dict.fromkeys(ip_pattern.findall(text)))[:50]

        # Detect scan tool
        scan_tool = "unknown"
        for tool in ["nessus", "qualys", "openvas", "nmap", "burp", "acunetix", "rapid7"]:
            if tool in text_lower:
                scan_tool = tool
                break

        return {
            "cve_count": len(cve_ids),
            "cve_ids": cve_ids[:30],
            "severity_counts": severity_counts,
            "unique_hosts": hosts[:20],
            "host_count": len(hosts),
            "scan_tool": scan_tool,
        }

    @staticmethod
    def _extract_config_metadata(text: str) -> Dict[str, Any]:
        """Detect secrets patterns and dangerous settings in config text."""
        secrets_patterns = [
            re.compile(r'(?:password|passwd|pwd)\s*[=:]\s*\S+', re.IGNORECASE),
            re.compile(r'(?:api[_-]?key|apikey)\s*[=:]\s*\S+', re.IGNORECASE),
            re.compile(r'(?:secret|token)\s*[=:]\s*\S+', re.IGNORECASE),
            re.compile(r'(?:connection[_-]?string)\s*[=:]\s*\S+', re.IGNORECASE),
        ]
        secrets_found = sum(len(p.findall(text)) for p in secrets_patterns)

        danger_settings = []
        danger_checks = [
            ("debug", r'\bdebug\s*[=:]\s*(?:true|1|on|yes)\b'),
            ("verbose_errors", r'\b(?:display_errors|show_errors)\s*[=:]\s*(?:true|1|on)\b'),
            ("open_cors", r'\b(?:allow_origin|cors)\s*[=:]\s*\*'),
            ("http_only", r'\bhttp://'),
        ]
        for label, pattern in danger_checks:
            if re.search(pattern, text, re.IGNORECASE):
                danger_settings.append(label)

        return {
            "secrets_found": secrets_found,
            "dangerous_settings": danger_settings,
        }

    @staticmethod
    def _extract_policy_metadata(text: str) -> Dict[str, Any]:
        """Extract framework references and section headings from policy text."""
        text_lower = text.lower()
        frameworks_detected = []
        framework_patterns = {
            "NIST CSF": r'nist\s+(?:csf|cybersecurity\s+framework)',
            "ISO 27001": r'iso\s*27001',
            "SOC 2": r'soc\s*2',
            "HIPAA": r'hipaa',
            "PCI DSS": r'pci[\s-]*dss',
            "GDPR": r'gdpr',
            "CIS Controls": r'cis\s+controls',
            "NIST 800-53": r'nist\s+800-53',
        }
        for name, pattern in framework_patterns.items():
            if re.search(pattern, text_lower):
                frameworks_detected.append(name)

        # Extract section headings (lines that look like headings)
        heading_pattern = re.compile(r'^(?:#{1,4}\s+|(?:\d+\.)+\s+)(.+)$', re.MULTILINE)
        headings = [m.group(1).strip() for m in heading_pattern.finditer(text)][:20]

        return {
            "frameworks_referenced": frameworks_detected,
            "section_headings": headings,
        }

    @staticmethod
    def _extract_architecture_metadata(text: str) -> Dict[str, Any]:
        """Extract technology names, components, and protocols from architecture text."""
        text_lower = text.lower()

        tech_keywords = {
            "aws": ["ec2", "s3", "rds", "lambda", "ecs", "eks", "vpc", "cloudfront", "api gateway"],
            "azure": ["vm", "blob", "cosmos", "aks", "app service", "front door"],
            "gcp": ["compute engine", "cloud storage", "bigquery", "gke", "cloud run"],
            "general": ["docker", "kubernetes", "nginx", "apache", "redis", "postgresql",
                       "mysql", "mongodb", "kafka", "rabbitmq", "elasticsearch"],
        }
        components_found = []
        for category, terms in tech_keywords.items():
            for term in terms:
                if term in text_lower:
                    components_found.append(term)

        protocols = []
        for proto in ["https", "http", "ssh", "rdp", "ftp", "sftp", "smtp", "dns", "tls", "ssl", "ipsec"]:
            if proto in text_lower:
                protocols.append(proto)

        return {
            "components_detected": components_found[:20],
            "protocols_mentioned": protocols,
        }
