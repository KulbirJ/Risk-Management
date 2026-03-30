"""Generate a high-level system design document for the Meridian HR Portal."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    return table


def build_document():
    doc = Document()

    # ── Title page ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("HIGH-LEVEL SYSTEM DESIGN DOCUMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Meridian HR Portal", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\n\nDocument Classification: INTERNAL — CONFIDENTIAL\n").bold = True
    meta.add_run(f"Version: 1.4  |  Date: {datetime.date.today().strftime('%B %d, %Y')}\n")
    meta.add_run("Author: Infrastructure Operations Team\n")
    meta.add_run("Owner: Corporate Services — HR Technology\n")
    meta.add_run("Status: AS-BUILT (reflects current production state)\n")

    doc.add_page_break()

    # ── Document Control ──
    add_heading_styled(doc, "1. Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"], [
        ["1.0", "2011-08-15", "Vendor (Apex Solutions Inc.)", "Initial architecture for go-live"],
        ["1.1", "2014-03-22", "J. Patel, Infra Ops", "Updated after server room move"],
        ["1.2", "2018-06-10", "J. Patel, Infra Ops", "Added payroll module components"],
        ["1.3", "2023-01-20", "M. Khan, Infra Ops", "Noted EOL status of major components"],
        ["1.4", datetime.date.today().strftime("%Y-%m-%d"), "M. Khan, Infra Ops", "Pre-assessment refresh"],
    ])

    # ── Executive Summary ──
    add_heading_styled(doc, "2. Executive Summary")
    doc.add_paragraph(
        "The Meridian HR Portal is a business-critical internal web application that provides "
        "payroll processing, benefits enrollment, employee records management, and leave tracking "
        "for approximately 1,200 employees across three regional offices (Toronto HQ, Ottawa, "
        "Calgary). The system was developed in 2011 by Apex Solutions Inc. (contract terminated 2016) "
        "and has operated with minimal changes since a 2018 payroll module addition."
    )
    doc.add_paragraph(
        "The application handles Protected Class B information including Social Insurance Numbers "
        "(SIN), banking details, salary records, and personal health data. Despite its criticality, "
        "the system runs entirely on end-of-life infrastructure with no active development support, "
        "no disaster recovery capability, and significant architectural security weaknesses documented "
        "in Section 8 of this document."
    )

    # ── System Overview ──
    add_heading_styled(doc, "3. System Overview")
    doc.add_heading("3.1 Purpose and Scope", level=2)
    doc.add_paragraph(
        "The Meridian HR Portal serves as the single system of record for all human resources "
        "data within the organization. Business processes supported include:"
    )
    bullets = [
        "Bi-weekly payroll processing and pay stub generation",
        "Annual benefits enrollment and mid-year life event changes",
        "Employee onboarding / offboarding record management",
        "Vacation and sick leave requests and approvals",
        "Departmental headcount and compensation reporting",
        "T4 / ROE generation for year-end and terminations",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.2 Users", level=2)
    add_table(doc, ["User Group", "Count", "Access Level", "Authentication"], [
        ["General Employees", "~1,150", "Self-service (own records only)", "Username / password"],
        ["HR Administrators", "12", "Full read/write to all employee data", "Username / password"],
        ["Payroll Officers", "6", "Payroll module + banking details", "Username / password"],
        ["IT Operations (Admins)", "2", "Server + database + /admin console", "RDP + /admin panel"],
        ["Finance (Read-only)", "8", "Compensation reports", "Username / password"],
    ])

    # ── Architecture ──
    add_heading_styled(doc, "4. Architecture")
    doc.add_heading("4.1 Deployment Architecture (Text Diagram)", level=2)
    diagram = doc.add_paragraph()
    diagram.style = doc.styles["No Spacing"]
    diagram_text = """
┌──────────────────────────────────────────────────────────────────────┐
│                        CORPORATE LAN (10.0.50.0/24)                  │
│                        *** FLAT NETWORK — NO SEGMENTATION ***        │
│                                                                      │
│   ┌──────────────┐        ┌────────────────────────────────┐         │
│   │  User         │        │  MERIDIAN-HR  (10.0.50.10)     │         │
│   │  Workstations │──HTTP──│  Windows Server 2008 R2 SP1    │         │
│   │  (~1,200)     │  :80   │                                │         │
│   │               │  :443  │  ┌─────────────────────────┐   │         │
│   └──────────────┘        │  │ Apache 2.2.34 (Win32)   │   │         │
│                            │  │ + mod_ssl (OpenSSL 1.0.2u│   │         │
│                            │  │ + PHP 5.6.40            │   │         │
│   ┌──────────────┐        │  └─────────┬───────────────┘   │         │
│   │  IT Admin     │──RDP───│             │                  │         │
│   │  Workstations │ :3389  │  ┌─────────▼───────────────┐   │         │
│   │  (2)          │        │  │ MySQL 5.5.62            │   │         │
│   └──────────────┘        │  │ Port 3306 (0.0.0.0)     │   │         │
│                            │  │ 42 tables, ~180K rows   │   │         │
│                            │  └─────────────────────────┘   │         │
│                            │                                │         │
│                            │  C:\\MeridianApp\\             │         │
│                            │    ├── htdocs/  (PHP source)   │         │
│                            │    ├── config.php (DB creds)   │         │
│                            │    ├── uploads/ (employee docs)│         │
│                            │    └── logs/                   │         │
│                            └────────────────────────────────┘         │
│                                         │                            │
│                                    FTP :21                           │
│                                         │                            │
│   ┌────────────────────────────────────▼──────────────────────┐      │
│   │  FTP-RELAY (10.0.50.25) — Ubuntu 16.04 / vsFTPd 3.0.2    │      │
│   │  Receives nightly payroll extracts (CSV) from MERIDIAN-HR │      │
│   │  Finance team picks up files manually each morning        │      │
│   └───────────────────────────────────────────────────────────┘      │
│                                                                      │
│   ┌───────────────────────────────────────────────────────────┐      │
│   │  NAS-BACKUP (10.0.50.30) — Synology DS920+               │      │
│   │  Nightly Windows Server Backup via SMBv1 share            │      │
│   │  \\\\NAS-BACKUP\\meridian-backup$                           │      │
│   │  Retention: 14 days rolling, no offsite copy              │      │
│   └───────────────────────────────────────────────────────────┘      │
│                                                                      │
│   ** No firewall between segments — all hosts on same VLAN **        │
│   ** No IDS/IPS — No WAF — No SIEM **                                │
└──────────────────────────────────────────────────────────────────────┘
"""
    run = diagram.add_run(diagram_text)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)

    doc.add_heading("4.2 Component Summary", level=2)
    add_table(doc, ["Component", "Technology", "Version", "EOL Status"], [
        ["Operating System", "Windows Server 2008 R2 SP1", "6.1.7601", "EOL Jan 2020"],
        ["Web Server", "Apache HTTP Server", "2.2.34", "EOL Dec 2017"],
        ["Application Runtime", "PHP", "5.6.40", "EOL Dec 2018"],
        ["Database", "MySQL Community", "5.5.62", "EOL Dec 2018"],
        ["TLS Library", "OpenSSL", "1.0.2u", "EOL Dec 2019"],
        ["FTP Server", "vsFTPd on Ubuntu 16.04", "3.0.2", "Ubuntu 16.04 EOL Apr 2021"],
        ["Backup Target", "Synology NAS (SMBv1)", "DSM 6.2", "SMBv1 deprecated"],
        ["Antivirus", "Symantec EP", "14.0 (defs stale 6mo+)", "Agent unmanaged"],
    ])

    # ── Authentication ──
    add_heading_styled(doc, "5. Authentication & Access Control")
    doc.add_paragraph(
        "The application implements its own authentication system entirely in PHP. "
        "There is no integration with Active Directory, LDAP, SSO, or any identity provider."
    )
    doc.add_heading("5.1 Authentication Mechanism", level=2)
    issues = [
        "Passwords stored as unsalted MD5 hashes in the users table",
        "No password complexity requirements enforced",
        "No account lockout after failed login attempts",
        "No multi-factor authentication (MFA)",
        "Session tokens passed as URL query parameters (?sid=...)",
        "Session timeout set to 24 hours (hardcoded)",
        "No CSRF protection on any form",
        "Remember-me cookie stored in plaintext with user ID",
    ]
    for i in issues:
        p = doc.add_paragraph(i, style="List Bullet")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_heading("5.2 Authorization", level=2)
    doc.add_paragraph(
        "Role-based access is implemented via a 'role' column in the users table with values: "
        "'employee', 'hr_admin', 'payroll', 'finance', 'it_admin'. Authorization checks are "
        "performed inconsistently — several admin-only pages only check for a valid session, "
        "not the user's role. The /admin panel uses a separate hardcoded credential (admin/admin123) "
        "that has not been changed since initial deployment."
    )

    # ── Data Flow ──
    add_heading_styled(doc, "6. Data Flow")
    doc.add_heading("6.1 Standard User Workflow", level=2)
    steps = [
        "Employee opens browser → http://meridian-hr.corp.local (HTTP, no TLS redirect)",
        "PHP login page → credentials submitted via POST over HTTP (plaintext on wire)",
        "PHP compares MD5(input) to stored hash → sets session ID in URL query string",
        "All subsequent requests carry ?sid=<token> → PHP reads session from /tmp/",
        "Data queries execute raw SQL (string concatenation) against MySQL",
        "Results rendered server-side in PHP → returned as HTML",
    ]
    for idx, s in enumerate(steps, 1):
        doc.add_paragraph(f"Step {idx}: {s}", style="List Number")

    doc.add_heading("6.2 Nightly Payroll Extract", level=2)
    steps2 = [
        "Scheduled Task (runs as svc_meridian) executes payroll_export.php at 01:00 EST",
        "PHP script queries all employee payroll records including SIN and banking data",
        "Output written to CSV: C:\\MeridianApp\\exports\\payroll_YYYYMMDD.csv",
        "Script connects to FTP-RELAY (10.0.50.25) via cleartext FTP and uploads CSV",
        "Finance team manually downloads file the next morning via FTP client",
        "CSV files on FTP-RELAY are never deleted (3+ years of payroll data accumulated)",
    ]
    for idx, s in enumerate(steps2, 1):
        doc.add_paragraph(f"Step {idx}: {s}", style="List Number")

    # ── Network Architecture ──
    add_heading_styled(doc, "7. Network Architecture")
    doc.add_paragraph(
        "All servers and workstations reside on a single flat VLAN (10.0.50.0/24). "
        "There is no network segmentation between the database server, web server, "
        "user workstations, or management interfaces. Key network architecture concerns:"
    )
    net_issues = [
        "No firewall between internal zones — all 1,200 workstations have direct access to MySQL (3306), RDP (3389), SMB (445), and FTP (21)",
        "No Web Application Firewall (WAF) in front of the application",
        "No Intrusion Detection/Prevention System (IDS/IPS)",
        "No network-level encryption — all internal traffic is plaintext",
        "DNS resolution via corporate DNS; no DNS logging or filtering",
        "Default gateway is a Cisco ASA 5505 (EOL firmware) providing NAT to the internet — application is NOT internet-facing but the server can reach the internet for NTP and (incorrectly) for Windows Update checks that always fail",
        "Wi-Fi guest network on the same physical switch, separated only by VLAN tagging (no ACLs between VLANs)",
    ]
    for n in net_issues:
        doc.add_paragraph(n, style="List Bullet")

    # ── Risk Summary ──
    add_heading_styled(doc, "8. Known Risks & Design Weaknesses")
    doc.add_paragraph(
        "The following risks have been identified through operational experience and the "
        "previous 2023 assessment. Items marked with ★ are new since the last review."
    )
    risks = [
        ["R-001", "Critical", "All infrastructure components are end-of-life with no security patches available"],
        ["R-002", "Critical", "SQL injection vulnerabilities throughout the application (no parameterized queries)"],
        ["R-003", "Critical", "Passwords stored as unsalted MD5 — trivially crackable"],
        ["R-004", "Critical", "Shared service account (svc_meridian) has Domain Admin privileges; credentials in plaintext config file"],
        ["R-005", "Critical", "No network segmentation — flat LAN exposes all services to all users"],
        ["R-006", "Critical", "EternalBlue (MS17-010) and BlueKeep (CVE-2019-0708) exploitable remotely — no patches available"],
        ["R-007", "High", "Sensitive data (SIN, banking) transferred via cleartext FTP nightly"],
        ["R-008", "High", "Self-signed TLS certificate expired Nov 2024; 1024-bit RSA key; SHA-1 signature"],
        ["R-009", "High", "Session tokens in URL query strings enable session hijacking via Referer leakage"],
        ["R-010", "High", "Admin console (/admin) uses default credentials (admin/admin123)"],
        ["R-011", "High", "No MFA on any account including admin and payroll"],
        ["R-012", "High", "No centralized logging or SIEM — incidents cannot be detected or investigated"],
        ["R-013", "High", "★ Antivirus definitions >6 months stale; agent not centrally managed"],
        ["R-014", "High", "★ Backups via SMBv1 to local NAS only; no offsite / no tested restore"],
        ["R-015", "Medium", "No change management process — patches applied ad-hoc without testing"],
        ["R-016", "Medium", "★ No disaster recovery plan; RPO/RTO undefined"],
        ["R-017", "Medium", "USB autorun enabled on server"],
        ["R-018", "Medium", "Directory listing enabled — config.php and database backups publicly browsable"],
        ["R-019", "Low", "★ phpMyAdmin accessible on port 8080 with MySQL root credentials"],
        ["R-020", "Low", "★ 3+ years of undeleted payroll CSVs on FTP relay server"],
    ]
    add_table(doc, ["Risk ID", "Severity", "Description"], risks)

    # ── Backup & DR ──
    add_heading_styled(doc, "9. Backup & Disaster Recovery")
    doc.add_paragraph(
        "Backups are performed nightly at 02:00 EST using Windows Server Backup. The backup "
        "target is a Synology DS920+ NAS (10.0.50.30) accessed via an SMBv1 share. Retention "
        "is 14 days rolling. There is no offsite backup copy, no cloud replication, and no "
        "tested restore procedure. The last successful restore test was never performed."
    )
    doc.add_paragraph(
        "There is no disaster recovery (DR) site or plan. The Business Continuity Plan "
        "defines a 4-hour RTO, but no technical capability exists to meet this target. "
        "If the physical server fails, estimated recovery time is 2-4 weeks (procure "
        "new hardware + rebuild from backup, if backup is intact)."
    )

    # ── Compliance ──
    add_heading_styled(doc, "10. Regulatory & Compliance Considerations")
    compliance = [
        "PIPEDA — Personal information of Canadian employees (SIN, health, banking) must be protected with appropriate safeguards. Current architecture does not meet reasonable security expectations.",
        "PCI DSS — Not directly in scope, but payroll banking data handling shares similar risk profile.",
        "CCCS ITSG-33 — As a federal contractor, the organization must align to ITSG-33 Medium integrity. Multiple controls are not met (AC-2, AC-17, AU-2, IA-2, SC-8, SC-13, SI-2, SI-3).",
        "SOC 2 Type II — The organization is pursuing SOC 2 certification. This system will fail Trust Service Criteria for Security (CC6, CC7, CC8) and Availability (A1).",
    ]
    for c in compliance:
        doc.add_paragraph(c, style="List Bullet")

    # ── Appendix ──
    add_heading_styled(doc, "Appendix A — Service Account Details")
    add_table(doc, ["Account", "Purpose", "Privileges", "Credential Storage"], [
        ["svc_meridian", "App-to-DB, FTP, scheduled tasks", "Domain Admin (!)", "Plaintext in config.php"],
        ["mysql root", "Database administration", "MySQL superuser", "Plaintext in config.php"],
        ["admin (web)", "/admin console", "Full app admin + DB query", "Hardcoded default (admin/admin123)"],
        ["ftp_payroll", "FTP file transfer", "Read/write on /payroll/", "Plaintext in payroll_export.php"],
    ])

    doc.add_paragraph("")
    add_heading_styled(doc, "Appendix B — Open Ports on MERIDIAN-HR (10.0.50.10)")
    add_table(doc, ["Port", "Protocol", "Service", "Notes"], [
        ["21", "TCP", "FTP (vsFTPd relay)", "Cleartext; used for payroll export"],
        ["80", "TCP", "HTTP (Apache 2.2.34)", "No redirect to HTTPS"],
        ["135", "TCP", "MSRPC", "Windows default"],
        ["139", "TCP", "NetBIOS Session", "SMBv1"],
        ["443", "TCP", "HTTPS (Apache + OpenSSL 1.0.2u)", "Self-signed expired cert"],
        ["445", "TCP", "SMB (SMBv1)", "EternalBlue vulnerable"],
        ["3306", "TCP", "MySQL 5.5.62", "Bound to 0.0.0.0; remote root enabled"],
        ["3389", "TCP", "RDP", "NLA disabled; BlueKeep vulnerable"],
        ["8080", "TCP", "phpMyAdmin", "MySQL root access via browser"],
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— END OF DOCUMENT —")
    run.bold = True
    run.font.size = Pt(14)

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = r"sample-data\Meridian_HR_Portal_System_Design.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
